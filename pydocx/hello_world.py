#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

import urllib2
import StringIO

if __name__ == '__main__':

    document = Document()
    h = document.add_heading(u'我是标题', 0)
    h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    document.add_paragraph(u'我是普通文档')
    document.add_paragraph(u'我是好看的文档', style='IntenseQuote')
    p = document.add_paragraph('')
    run1 = p.add_run(u'句子一')
    run1.bold = True
    run2 = p.add_run(u'句子二')
    run2.italic = True

    t = document.add_table(rows=2, cols=2)
    t.cell(0, 0).text = u'单元1'
    t.cell(0, 1).text = u'单元2'
    t.cell(1, 0).text = u'单元3'
    t.cell(1, 1).text = u'单元4'

    t = document.add_table(rows=2, cols=2)
    merge_cell = t.cell(0, 0).merge(t.cell(0, 1))
    merge_cell.text = u'合并单元格'

    t = document.add_table(rows=2, cols=2)
    t.style = 'Table Grid'

    document.add_picture("./media/logo_for_1.png")
    img_url = "http://zeromq.wdfiles.com/local--files/admin:css/logo.gif"
    image_from_url = urllib2.urlopen(img_url)
    io_url = StringIO.StringIO()
    io_url.write(image_from_url.read())
    io_url.seek(0)
    document.add_picture(io_url)

    # print "console------"
    # print len(document.styles)
    # for s in document.styles:
    #     print s.name
    
    new_section = document.add_section(WD_SECTION.ODD_PAGE)
    
    document.save('./store/foo.docx')


